# services/template_apply.py
from __future__ import annotations

from pathlib import Path
from typing import Dict, Optional, List
import os
import re
import shutil
import zipfile
import xml.etree.ElementTree as ET

from .db import get_project
from .logo_store import list_logos
from .templates_store import resolve_abs_path  # same package

# ================== Behaviour flags ==========================================
# Excel handling
USE_XLWINGS_FIRST = True           # preserve CF/DataValidation/etc.
ALLOW_OPENPYXL_FALLBACK = False    # only if you're ok with some feature loss

# Word handling
WORD_UNLINK_FIELDS_VIA_COM = False  # set True to flatten fields via COM first
WORD_ENABLE_XML_PASS = True         # replace inside shapes/text boxes via XML
WORD_DEBUG = False                  # verbose docx debug

# ================== Utilities / logging ======================================
_CAT_TO_WORD = {
    "document":     "Documents",
    "schedule":     "Schedules",
    "drawing":      "Drawings",
    "rfi":          "RFI",
    "calculation":  "Model & Calc",
}


def _dbg(msg: str) -> None:
    try:
        print(f"[template_apply] {msg}", flush=True)
    except Exception:
        pass

def _wdbg(msg: str) -> None:
    if WORD_DEBUG:
        _dbg(f"[word] {msg}")

def _find_or_create_category_dir(register_path: Path, category: str) -> Path:
    """
    Given the path to the project DB (which lives under “…/1 Doc Control/…”),
    return the destination folder (e.g., “…/6 Documents”) for the given category.

    Rules:
      • Go to the folder ABOVE “Doc Control”.
      • Match target folders by name only (ignore any leading number and spaces).
      • If the folder does not exist, create it under the project root.
    """
    rp = Path(register_path).resolve()

    def norm(name: str) -> str:
        # strip leading digits + whitespace, compare case-insensitively
        return re.sub(r"^\s*\d+\s*", "", name).strip().lower()

    # Find “…/1 Doc Control/” up the tree, then step up to the project root
    cur = rp.parent
    base_dir = None
    for _ in range(4):  # climb a few levels just in case (handles “…/1 Doc Control/.docutrans/DB.db” too)
        if norm(cur.name) in {"doc control", "doccontrol"}:
            base_dir = cur.parent
            break
        cur = cur.parent
    if base_dir is None:
        # Fallback: assume DB is directly under “…/1 Doc Control/DB.db”
        base_dir = rp.parent.parent

    want = _CAT_TO_WORD.get((category or "").lower(), "Documents")
    want_norm = norm(want)

    # Look for an existing sibling like “3 Drawings”, “4 Schedules”, “5 Model & Calc”, etc.
    try:
        for p in base_dir.iterdir():
            if p.is_dir() and norm(p.name) == want_norm:
                return p
    except Exception:
        pass

    # Create if not found
    out = base_dir / want
    out.mkdir(parents=True, exist_ok=True)
    return out

# ================== Excel: xlwings (preferred) ===============================
def _apply_excel_with_xlwings(dest_path: Path, doc_id: str, project: Dict[str, str], logos: List[Path]) -> None:
    _dbg("trying xlwings route...")
    import xlwings as xw

    app = None
    book = None
    try:
        app = xw.App(visible=False, add_book=False)
        book = xw.Book(str(dest_path))

        # Sheet selection
        sheet_names = [sht.name for sht in book.sheets]
        sht = book.sheets["Cover Sheet"] if "Cover Sheet" in sheet_names else book.sheets[0]

        # Fill cells
        sht.range("I4").value  = doc_id
        sht.range("I6").value  = project.get("project_code") or ""
        sht.range("I8").value  = project.get("client_company") or project.get("client_reference") or ""
        sht.range("I10").value = project.get("end_user") or ""

        # Clear A7 text
        try:
            sht.range("A7").value = None
        except Exception:
            pass

        # --- Logo layout in the A7 zone (keeps aspect, centers in slots) ---
        try:
            zone = sht.range("A7").merge_area
        except Exception:
            zone = sht.range("A7")

        # Remove our previous pictures if re-run
        try:
            for pic in list(sht.pictures):
                if str(getattr(pic, "name", "")).startswith("ClientLogo_"):
                    pic.delete()
        except Exception:
            pass

        max_logos = min(len(logos), 3)
        if max_logos > 0:
            gutter = 6.0
            slot_w = (zone.width - gutter * (max_logos - 1)) / max(1, max_logos)
            slot_h = zone.height

            for i, p in enumerate(logos[:max_logos]):
                left_slot = zone.left + i * (slot_w + gutter)
                top_slot = zone.top
                pic = sht.pictures.add(str(p), name=f"ClientLogo_{i+1}", left=left_slot, top=top_slot)
                w0, h0 = float(pic.width), float(pic.height)
                if w0 > 0 and h0 > 0:
                    scale = min(slot_w / w0, slot_h / h0, 1.0)
                    pic.width = w0 * scale
                    pic.height = h0 * scale
                    pic.left = left_slot + (slot_w - pic.width) / 2.0
                    pic.top = top_slot + (slot_h - pic.height) / 2.0

        book.save()
        _dbg("xlwings apply complete (saved by Excel).")
    finally:
        try:
            if book is not None:
                book.close()
        except Exception:
            pass
        try:
            if app is not None:
                app.quit()
        except Exception:
            pass

# ================== Excel: openpyxl fallback (feature loss possible) =========
def _apply_excel_with_openpyxl(dest_path: Path, doc_id: str, project: Dict[str, str], logos: List[Path]) -> None:
    _dbg("using openpyxl fallback (may remove advanced formatting/validation).")
    try:
        from openpyxl import load_workbook
        from openpyxl.drawing.image import Image as XLImage
    except Exception as e:
        _dbg(f"openpyxl not available: {e}")
        return

    wb = load_workbook(str(dest_path))
    ws = wb["Cover Sheet"] if "Cover Sheet" in wb.sheetnames else wb.active

    ws["I4"]  = doc_id
    ws["I6"]  = project.get("project_code") or ""
    ws["I8"]  = project.get("client_company") or project.get("client_reference") or ""
    ws["I10"] = project.get("end_user") or ""

    try:
        ws["A7"].value = ""
    except Exception:
        pass

    anchors = ["A7", "F7", "K7"]
    for i, p in enumerate(logos[:3]):
        try:
            img = XLImage(str(p))
            ws.add_image(img, anchors[i if i < len(anchors) else -1])
        except Exception:
            continue

    wb.save(str(dest_path))
    _dbg("openpyxl apply complete (file saved).")

# ================== Word: mapping, python-docx and XML pass ===================
def _word_mapping_from_project(project: Dict[str, str], payload: Dict[str, str]) -> Dict[str, str]:
    doc_id = (payload.get("doc_id") or "").strip()
    revision = (payload.get("revision") or payload.get("template_revision") or "").strip()
    description = (payload.get("description") or payload.get("template_description") or "").strip()
    return {
        "<<Project Name>>": project.get("project_name") or "",
        "<<Project No>>":   project.get("project_code") or "",
        "<<Client>>":       project.get("client_company") or project.get("client_reference") or "",
        "<<End User>>":     project.get("end_user") or "",
        "<<Doc ID>>":       doc_id,
        "<<Document ID>>":  doc_id,
        "<<Document No>>":  doc_id,
        "<<Revision>>":     revision,
        "<<Description>>":  description,
        "<<Client Reference>>": project.get("client_reference") or "",
        "<<Client Contact>>":   project.get("client_contact") or "",
    }

def _replace_in_paragraph(paragraph, mapping: Dict[str, str]) -> bool:
    runs = paragraph.runs
    if not runs:
        return False
    full_text = "".join(r.text or "" for r in runs)
    replaced = full_text
    for k, v in mapping.items():
        if k in replaced:
            replaced = replaced.replace(k, v)
    if replaced == full_text:
        return False

    # Keep first run style
    first_style = runs[0].style if runs[0].style else None
    # Remove all runs except first
    for i in range(len(runs) - 1, 0, -1):
        r = runs[i]
        r._element.getparent().remove(r._element)
    runs[0].text = replaced
    if first_style:
        runs[0].style = first_style
    return True

def _replace_in_table(table, mapping: Dict[str, str]) -> int:
    count = 0
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                if _replace_in_paragraph(p, mapping):
                    count += 1
            for t2 in cell.tables:  # nested
                count += _replace_in_table(t2, mapping)
    return count

def _apply_word_with_python_docx(dest_path: Path, project: Dict[str, str], payload: Dict[str, str]) -> int:
    try:
        from docx import Document
    except Exception as e:
        _dbg(f"python-docx not available: {e}")
        return 0

    mapping = _word_mapping_from_project(project, payload)
    _wdbg(f"mapping keys={list(mapping.keys())}")

    doc = Document(str(dest_path))
    total = 0

    # body paragraphs
    for p in doc.paragraphs:
        if _replace_in_paragraph(p, mapping):
            total += 1

    # body tables
    for t in doc.tables:
        total += _replace_in_table(t, mapping)

    # headers/footers
    try:
        for sec in doc.sections:
            try:
                for p in sec.header.paragraphs:
                    if _replace_in_paragraph(p, mapping):
                        total += 1
                for t in sec.header.tables:
                    total += _replace_in_table(t, mapping)
            except Exception:
                pass
            try:
                for p in sec.footer.paragraphs:
                    if _replace_in_paragraph(p, mapping):
                        total += 1
                for t in sec.footer.tables:
                    total += _replace_in_table(t, mapping)
            except Exception:
                pass
    except Exception:
        pass

    doc.save(str(dest_path))
    _dbg(f"word python-docx replace complete ({total} paragraphs/cells updated).")
    return total

def _apply_word_xml_replace(dest_path: Path, mapping: Dict[str, str]) -> int:
    """
    Safe, fast DOCX XML replace:
      - Processes each <w:p> independently (body, headers, footers, text boxes)
      - Replaces in descendant <w:t> preserving run boundaries
      - Uses lxml when available (recommended), else xml.etree
      - Returns count of ZIP parts changed
    """
    try:
        # Prefer lxml for better namespace handling & stable serialization
        import lxml.etree as ETmod
        use_lxml = True
    except Exception:
        import xml.etree.ElementTree as ETmod  # type: ignore
        use_lxml = False

    W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    WPS_NS = "http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
    NSMAP = {"w": W_NS, "wp": WP_NS, "wps": WPS_NS}

    SEP = "\uF000"  # sentinel used within a single <w:p>

    def _serialize(root):
        if use_lxml:
            return ETmod.tostring(
                root, encoding="utf-8", xml_declaration=True, standalone=False
            )
        # stdlib ElementTree
        return ETmod.tostring(root, encoding="utf-8")

    def _patch_paragraph(p) -> bool:
        """
        Replace tokens across the concatenated text of all descendant <w:t> of this <w:p>.
        If token spans multiple runs, it still gets replaced. Returns True if changed.
        """
        ts = list(p.findall(".//{"+W_NS+"}t"))
        if not ts:
            return False

        orig_pieces = [t.text or "" for t in ts]
        joined = SEP.join(orig_pieces)

        new = joined
        for k, v in mapping.items():
            if k and (k in new):
                new = new.replace(k, v)

        if new == joined:
            return False

        new_pieces = new.split(SEP)
        if len(new_pieces) == len(orig_pieces):
            for t, seg in zip(ts, new_pieces):
                t.text = seg
            return True

        # Fallback: per-node replace (handles odd segment counts safely)
        changed = False
        for t in ts:
            txt = t.text or ""
            rep = txt
            for k, v in mapping.items():
                if k in rep:
                    rep = rep.replace(k, v)
            if rep != txt:
                t.text = rep
                changed = True
        return changed

    def patch_xml(data: bytes) -> tuple[bytes, bool]:
        try:
            if use_lxml:
                parser = ETmod.XMLParser(remove_blank_text=False, recover=False)
                root = ETmod.fromstring(data, parser)
            else:
                root = ETmod.fromstring(data)
        except Exception:
            # If XML can’t be parsed, don’t touch this part
            return data, False

        # Process every paragraph in this part, including inside text boxes
        changed = False
        for p in root.findall(".//{"+W_NS+"}p"):
            if _patch_paragraph(p):
                changed = True

        if not changed:
            return data, False

        return _serialize(root), True

    # ---- zip roundtrip, modifying only the Word parts that contain text ----
    changed_parts = 0
    src = str(dest_path)
    tmp = src + ".tmp"
    backup = src + ".bak"  # keep one-time backup to help debug unreadable files

    with zipfile.ZipFile(src, "r") as zin, zipfile.ZipFile(tmp, "w", compression=zipfile.ZIP_DEFLATED) as zout:
        for info in zin.infolist():
            data = zin.read(info.filename)
            needs_patch = (
                info.filename == "word/document.xml"
                or info.filename.startswith("word/header")
                or info.filename.startswith("word/footer")
            )
            if needs_patch:
                data2, changed = patch_xml(data)
                if changed:
                    changed_parts += 1
                    data = data2
            # write with a fresh ZipInfo to avoid odd flags from original entries
            zi = zipfile.ZipInfo(filename=info.filename, date_time=info.date_time)
            zi.compress_type = zipfile.ZIP_DEFLATED
            zout.writestr(zi, data)

    # keep a backup only once (if not already present)
    try:
        if not os.path.exists(backup):
            shutil.copy2(src, backup)
    except Exception:
        pass
    os.replace(tmp, src)
    return changed_parts


def _unlink_word_fields(dest_path: Path) -> None:
    """
    Optional: fast COM pass to unlink all fields (turn into text)
    in body + headers/footers. Disabled by default.
    """
    import pythoncom
    import win32com.client as win32

    pythoncom.CoInitialize()
    app = None
    doc = None
    try:
        app = win32.Dispatch("Word.Application")
        app.Visible = False
        app.DisplayAlerts = 0  # wdAlertsNone
        doc = app.Documents.Open(str(dest_path))

        # body
        try:
            for f in list(doc.Fields):
                f.Unlink()
        except Exception:
            pass

        # headers/footers
        try:
            wdHeaderFooterPrimary = 1
            wdHeaderFooterFirst   = 2
            wdHeaderFooterEven    = 3
            for sec in doc.Sections:
                for group in (sec.Headers, sec.Footers):
                    for idx in (wdHeaderFooterPrimary, wdHeaderFooterFirst, wdHeaderFooterEven):
                        try:
                            for f in list(group(idx).Range.Fields):
                                f.Unlink()
                        except Exception:
                            pass
        except Exception:
            pass

        doc.Save()
        _dbg("word fields unlinked (flattened to text).")
    finally:
        try:
            if doc is not None:
                doc.Close(False)
        except Exception:
            pass
        try:
            if app is not None:
                app.Quit()
        except Exception:
            pass

# ================== Public entrypoint =========================================
def apply_template_for_new_doc(register_path: Path, payload: Dict[str, str]) -> Optional[Path]:
    _dbg(
        f"called use_template={bool(payload.get('use_template'))} "
        f"doc_id='{payload.get('doc_id','')}' "
        f"category='{payload.get('template_category','')}' "
        f"kind='{payload.get('template_kind','')}'"
    )

    if not payload or not payload.get("use_template"):
        return None

    t_cat  = (payload.get("template_category") or "document").lower()
    t_kind = (payload.get("template_kind") or "").lower()  # "excel" | "word"
    doc_id = (payload.get("doc_id") or "").strip()
    if not doc_id:
        _dbg("no doc_id supplied; aborting.")
        return None

    # Resolve source
    src: Optional[Path] = None
    if payload.get("template_abspath"):
        src = Path(payload["template_abspath"])
        _dbg(f"resolved from abspath: {src}")
    elif payload.get("template_relpath"):
        try:
            src = resolve_abs_path({"relpath": payload["template_relpath"]})
            _dbg(f"resolved from relpath via resolve_abs_path: {src}")
        except Exception as e:
            _dbg(f"resolve_abs_path error: {e}")
            src = None
    elif payload.get("template_path"):
        src = Path(payload["template_path"])
        _dbg(f"resolved from legacy template_path: {src}")

    if not src:
        _dbg("no source path resolved; aborting.")
        return None
    if not src.exists():
        _dbg(f"source path does not exist: {src}")
        return None

    dest_dir = _find_or_create_category_dir(Path(register_path), t_cat)
    _dbg(f"dest_dir: {dest_dir}")
    ext = src.suffix
    dest_path = _unique_path(dest_dir, doc_id, ext)
    _dbg(f"dest_path candidate: {dest_path}")

    # Copy source before modifying
    shutil.copy2(str(src), str(dest_path))
    _dbg(f"copied template to: {dest_path}")

    # Excel
    if t_kind == "excel" or src.suffix.lower() in (".xlsx", ".xlsm", ".xls"):
        proj = get_project(Path(register_path)) or {}
        logos = list_logos(Path(register_path))
        _dbg(
            "project fields: "
            f"code='{proj.get('project_code','')}', "
            f"client_company='{proj.get('client_company','')}', "
            f"client_ref='{proj.get('client_reference','')}', "
            f"end_user='{proj.get('end_user','')}'"
        )
        _dbg(f"logos found: {len(logos)} -> {[p.name for p in logos]}")

        if USE_XLWINGS_FIRST:
            try:
                import xlwings  # noqa: F401
                _apply_excel_with_xlwings(dest_path, doc_id, proj, logos)
                return dest_path
            except ImportError:
                _dbg("xlwings not installed.")
            except Exception as e:
                _dbg(f"xlwings error: {e}")

        if ALLOW_OPENPYXL_FALLBACK:
            _apply_excel_with_openpyxl(dest_path, doc_id, proj, logos)
            return dest_path

        _dbg("Skipped openpyxl fallback to preserve advanced features. File left copied but unfilled.")
        return dest_path

    # Word
    if t_kind == "word" or src.suffix.lower() in (".docx", ".docm", ".doc"):
        proj = get_project(Path(register_path)) or {}
        try:
            if WORD_UNLINK_FIELDS_VIA_COM:
                _unlink_word_fields(dest_path)  # flatten fields to text (optional)

            n1 = _apply_word_with_python_docx(dest_path, proj, payload)
            n2 = 0
            if WORD_ENABLE_XML_PASS:
                mapping = _word_mapping_from_project(proj, payload)
                n2 = _apply_word_xml_replace(dest_path, mapping)
                _dbg(f"word xml replace complete ({n2} parts updated).")

            if n1 == 0 and n2 == 0:
                _dbg("no word replacements were applied; placeholders may not be present.")
            return dest_path
        except Exception as e:
            _dbg(f"Word replace error: {e}")
            return dest_path

    # Fallback
    _dbg(f"fallback copy-only for unknown kind/ext: {src.suffix}")
    return dest_path


# --- replace your _unique_path with this version ---
from pathlib import Path

def _unique_path(dest_dir: Path, stem_or_name: str, ext: str = "") -> Path:
    """
    Return a unique file path in dest_dir.
    - If stem_or_name already has a suffix and ext is empty, keep that suffix.
    - If ext is provided, use it (with a leading dot if missing).
    - If the target exists, append ' (2)', ' (3)', ... before the suffix.
    """
    dest_dir = Path(dest_dir)
    dest_dir.mkdir(parents=True, exist_ok=True)

    base = (stem_or_name or "").strip() or "New Document"
    base = Path(base).name  # basename only

    # Decide suffix
    name_suffix = Path(base).suffix
    suffix = (ext or name_suffix or "")
    if suffix and not suffix.startswith("."):
        suffix = "." + suffix

    # Decide stem
    stem = Path(base).stem if suffix else base

    candidate = dest_dir / f"{stem}{suffix}"
    if not candidate.exists():
        return candidate

    i = 2
    while True:
        cand = dest_dir / f"{stem} ({i}){suffix}"
        if not cand.exists():
            return cand
        i += 1
